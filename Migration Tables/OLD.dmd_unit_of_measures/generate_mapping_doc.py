#!/usr/bin/env python3
"""Generate a Word (.docx) migration mapping document.

Creates `MIGRATION_MAPPING.docx` in the same folder.
"""
from docx import Document
from docx.shared import Inches
import os

OUT_NAME = "MIGRATION_MAPPING.docx"

ROWS = [
    ("id", "int8 NOT NULL", "`cd`", "Uses the source unit-of-measure code from MySQL `dmd_unit_of_measures` as the target Postgres ID."),
    ("description", "text NOT NULL", "`desc`", "Mapped from MySQL `desc` to the new Postgres description field."),
    ("created_at", "timestamptz NOT NULL", "created_at", "Uses source created_at when available; otherwise the target default now()."),
    ("updated_at", "timestamptz NOT NULL", "updated_at", "Uses source updated_at when available; otherwise the target default now()."),
]


def make_doc(out_path: str):
    doc = Document()
    doc.add_heading('Migration mapping: MySQL qitech.dmd_unit_of_measures -> Postgres public.dmd_lookup_units_of_measure', level=1)
    doc.add_paragraph('This document lists each target column, the Postgres type, the source column in MySQL, and any transformations or notes.')

    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Target column'
    hdr_cells[1].text = 'Postgres type'
    hdr_cells[2].text = 'Source (MySQL qitech.dmd_unit_of_measures)'
    hdr_cells[3].text = 'Transformation / Notes'

    for target, ptype, source, notes in ROWS:
        row_cells = table.add_row().cells
        row_cells[0].text = target
        row_cells[1].text = ptype
        row_cells[2].text = source
        row_cells[3].text = notes

    doc.add_heading('Additional notes', level=2)
    doc.add_paragraph(' - The script uses `cd` from MySQL as the Postgres target `id` value and converts it to an integer.')
    doc.add_paragraph(' - `desc` is mapped to `description`; `cd_prev` and `cd_date` are not migrated.')
    doc.add_paragraph(' - `created_at` and `updated_at` use the source values when provided; otherwise the current UTC timestamp is used.')
    doc.add_paragraph(' - Rows with duplicate `id` values are skipped in the target table.')

    doc.add_heading('Example transformed JSON for one row', level=2)
    doc.add_paragraph('''{
  "id": 1001,
  "description": "Kilogram",
  "created_at": "2024-01-01T12:00:00Z",
  "updated_at": "2024-01-02T12:00:00Z"
}''')

    doc.add_heading('Recommended actions before running import', level=2)
    doc.add_paragraph(' - Verify APID completeness (no NULLs) or modify the script to skip/fallback.')
    doc.add_paragraph(' - Backup the target Postgres table or run in staging first.')
    doc.add_paragraph(' - Run with --dry-run to review counts and sample transformed rows.')

    doc.save(out_path)


if __name__ == '__main__':
    here = os.path.dirname(__file__)
    out = os.path.join(here, OUT_NAME)
    make_doc(out)
    print(f'Wrote {out}')
